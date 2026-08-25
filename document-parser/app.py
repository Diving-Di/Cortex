import io
import os
import subprocess
import tempfile
import zipfile
from pathlib import Path

import fitz
import pytesseract
from docx import Document
from fastapi import FastAPI, Header, HTTPException, Request
from PIL import Image

app = FastAPI(docs_url=None, redoc_url=None, openapi_url=None)
MAX_BYTES = int(os.getenv("PARSER_MAX_BYTES", str(64 * 1024 * 1024)))
MAX_PAGES = int(os.getenv("PARSER_MAX_PAGES", "500"))
MAX_PIXELS = int(os.getenv("PARSER_MAX_PIXELS", "80000000"))
VERSION = "cortex-parser/1"
Image.MAX_IMAGE_PIXELS = MAX_PIXELS


def block(text, block_type="paragraph", page=None, image_index=None, confidence=None):
    return {"block_type": block_type, "text": text.strip(), "heading_path": [], "page": page,
            "image_index": image_index, "source_span": "", "parser_version": VERSION,
            "ocr_confidence": confidence}


def ocr_image(image, page=None, image_index=None):
    data = pytesseract.image_to_data(image, lang="chi_sim+eng", output_type=pytesseract.Output.DICT)
    words, scores = [], []
    for text, conf in zip(data["text"], data["conf"]):
        text = text.strip()
        try:
            score = float(conf)
        except (TypeError, ValueError):
            score = -1
        if text and score >= 0:
            words.append(text)
            scores.append(score / 100)
    if not words:
        raise HTTPException(424, "OCR produced no text")
    return block(" ".join(words), "ocr", page, image_index, sum(scores) / len(scores))


def parse_pdf(raw):
    try:
        document = fitz.open(stream=raw, filetype="pdf")
        if document.needs_pass:
            raise HTTPException(422, "encrypted document")
        if document.page_count > MAX_PAGES:
            raise HTTPException(413, "too many pages")
        result = []
        for number, page in enumerate(document, 1):
            text = page.get_text("text").strip()
            if text:
                result.append(block(text, "paragraph", number))
            else:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                result.append(ocr_image(Image.open(io.BytesIO(pix.tobytes("png"))), page=number))
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, "unsafe or corrupt PDF") from exc


def parse_docx(raw):
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            if any(name.lower().endswith("vbaproject.bin") for name in archive.namelist()):
                raise HTTPException(422, "macros are not accepted")
            expanded = sum(item.file_size for item in archive.infolist())
            if expanded > MAX_BYTES * 10:
                raise HTTPException(413, "expanded document too large")
        document = Document(io.BytesIO(raw))
        result = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                kind = "heading" if paragraph.style and paragraph.style.name.startswith("Heading") else "paragraph"
                result.append(block(paragraph.text, kind))
        for table in document.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            if any(rows):
                result.append(block("\n".join(rows), "table"))
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, "unsafe or corrupt DOCX") from exc


def parse_doc(raw):
    with tempfile.NamedTemporaryFile(suffix=".doc") as source:
        source.write(raw)
        source.flush()
        try:
            proc = subprocess.run(["antiword", source.name], capture_output=True, timeout=60, check=False)
        except subprocess.TimeoutExpired as exc:
            raise HTTPException(422, "DOC parse timed out") from exc
    if proc.returncode != 0:
        raise HTTPException(422, "unsafe or corrupt DOC")
    text = proc.stdout.decode("utf-8", errors="replace").strip()
    return [block(text)] if text else []


@app.get("/healthz")
def healthz():
    return {"status": "ok"}


@app.post("/v1/parse")
async def parse(request: Request, x_filename: str = Header(...)):
    filename = Path(x_filename).name
    ext = Path(filename).suffix.lower()
    raw = await request.body()
    if not raw or len(raw) > MAX_BYTES:
        raise HTTPException(413, "document too large")
    if ext == ".pdf" and raw.startswith(b"%PDF-"):
        blocks = parse_pdf(raw)
    elif ext == ".docx" and raw.startswith(b"PK"):
        blocks = parse_docx(raw)
    elif ext == ".doc" and raw.startswith(bytes.fromhex("d0cf11e0a1b11ae1")):
        blocks = parse_doc(raw)
    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        try:
            image = Image.open(io.BytesIO(raw)); image.verify()
            image = Image.open(io.BytesIO(raw)).convert("RGB")
            blocks = [ocr_image(image, image_index=1)]
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(400, "file type mismatch") from exc
    else:
        raise HTTPException(400, "file type mismatch")
    blocks = [item for item in blocks if item["text"]]
    if not blocks:
        raise HTTPException(422, "document contains no text")
    return {"blocks": blocks}
