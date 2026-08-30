import io
import os
import shutil
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app import parse_docx, parse_pdf, ocr_image


pdf = fitz.open()
page = pdf.new_page()
page.insert_text((72, 72), "Cortex PDF acceptance text")
assert "acceptance" in parse_pdf(pdf.tobytes())[0]["text"]

word = Document()
word.add_heading("Cortex Word", 1)
word.add_paragraph("DOCX acceptance text")
buffer = io.BytesIO()
word.save(buffer)
assert any("acceptance" in item["text"] for item in parse_docx(buffer.getvalue()))

if shutil.which("tesseract"):
    image = Image.new("RGB", (1000, 240), "white")
    candidates = [Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("C:/Windows/Fonts/arial.ttf")]
    font_path = next((path for path in candidates if path.exists()), None)
    font = ImageFont.truetype(str(font_path), 72) if font_path else ImageFont.load_default(size=72)
    ImageDraw.Draw(image).text((40, 70), "Cortex OCR 2026", fill="black", font=font)
    assert "2026" in ocr_image(image, image_index=1)["text"]
    ocr_status = "pass"
elif os.getenv("CI"):
    raise RuntimeError("tesseract is required in CI")
else:
    ocr_status = "skipped (tesseract unavailable)"

print(f"pdf=pass docx=pass image_ocr={ocr_status}")
