"""Generate deterministic, synthetic TXT/PDF/DOCX knowledge fixtures."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1] / "testdata" / "knowledge"


def fixture_text(chinese_phrase: str, english_phrase: str) -> str:
    return (
        f"苍穹计划的发布口令是{chinese_phrase}。这是完全虚构的自动化测试信息。\n"
        f"Project Sky uses the fictional release phrase {english_phrase}.\n"
        "混合说明：只有当前租户上传的资料可以被检索和引用。"
    )


def set_run_font(run, name: str = "Calibri", size: int = 11) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def generate_txt() -> None:
    for tenant, chinese, english in (
        ("bilingual-sample", "青竹七号", "Bamboo Seven"),
        ("tenant-a", "青竹七号", "Bamboo Seven"),
        ("tenant-b", "琥珀九号", "Amber Nine"),
    ):
        (ROOT / f"{tenant}.txt").write_text(
            fixture_text(chinese, english) + "\n", encoding="utf-8"
        )


def generate_pdf(tenant: str, launch_date: str, launch_date_zh: str) -> None:
    font_candidates = (
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    )
    font_path = next((path for path in font_candidates if path.exists()), None)
    if font_path is None:
        raise RuntimeError("A Unicode CJK TrueType font is required to generate PDF fixtures")
    if "CortexFixture" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("CortexFixture", str(font_path)))
    output = ROOT / f"{tenant}.pdf"
    document = canvas.Canvas(str(output), pagesize=letter)
    width, height = letter
    for page, lines in enumerate(
        [
            [
                "Cortex Knowledge Fixture",
                "中文知识样本",
                f"苍穹计划的虚构演练日期是{launch_date_zh}。",
            ],
            [
                "Cortex Knowledge Fixture",
                "English knowledge sample",
                f"Project Sky uses the fictional exercise date {launch_date}.",
            ],
            [
                "Cortex Knowledge Fixture",
                "中英混合样本",
                f"Only the current tenant may retrieve the date {launch_date}.",
            ],
        ],
        start=1,
    ):
        document.setFont("CortexFixture", 15)
        document.drawString(72, height - 72, lines[0])
        document.setFont("CortexFixture", 12)
        y = height - 120
        for line in lines[1:]:
            document.drawString(72, y, line)
            y -= 28
        document.drawCentredString(width / 2, 36, f"Page {page}")
        document.showPage()
    document.save()


def generate_docx(tenant: str, owner_zh: str, owner_en: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    document.add_heading("Cortex 知识库自动化样本", level=1)
    paragraph = document.add_paragraph()
    set_run_font(
        paragraph.add_run(
            f"苍穹计划的虚构负责团队是{owner_zh}。\n"
            f"The fictional owner of Project Sky is {owner_en}.\n"
            "混合说明：只有当前租户上传的资料可以被检索和引用。"
        )
    )
    document.add_heading("结构化数据", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.5)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "虚构值"
    for label, value in (
        ("项目", "苍穹计划 / Project Sky"),
        ("负责团队", f"{owner_zh} / {owner_en}"),
        ("用途", "双租户隔离与引用验收"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.save(ROOT / f"{tenant}.docx")


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    generate_txt()
    for tenant, launch_date, launch_date_zh, owner_zh, owner_en in (
        ("bilingual-sample", "2042-04-17", "2042年4月17日", "极光实验室", "Aurora Lab"),
        ("tenant-a", "2042-04-17", "2042年4月17日", "极光实验室", "Aurora Lab"),
        ("tenant-b", "2049-09-23", "2049年9月23日", "暮光工作室", "Twilight Studio"),
    ):
        generate_pdf(tenant, launch_date, launch_date_zh)
        generate_docx(tenant, owner_zh, owner_en)


if __name__ == "__main__":
    main()
